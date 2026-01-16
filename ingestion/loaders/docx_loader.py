"""Word document (DOCX) loader."""

from pathlib import Path

import structlog
from docx import Document as DocxDocument

from ingestion.loaders.base import BaseLoader
from schemas import Document

logger = structlog.get_logger(__name__)


class DocxLoader(BaseLoader):
    """Loader for Word documents (.docx)."""

    def supports(self, source: str | Path) -> bool:
        """Check if source is a DOCX file."""
        path = Path(source)
        return path.suffix.lower() == ".docx"

    async def load(self, source: str | Path) -> Document:
        """Load text content from DOCX.

        Args:
            source: Path to DOCX file.

        Returns:
            Document with extracted text content.
        """
        path = Path(source)
        logger.info("Loading DOCX file", path=str(path))

        doc = DocxDocument(path)
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        return Document(
            source=str(path),
            content=content,
            metadata={
                "file_type": "docx",
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "paragraph_count": len(paragraphs),
            },
        )
